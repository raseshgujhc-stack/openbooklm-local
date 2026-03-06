"""
HTTP API routes for STT file transcription and capability discovery.
"""

from datetime import datetime
import tempfile
from pathlib import Path
import io
import re
from typing import List, Optional, Tuple

from fastapi import APIRouter, File, HTTPException, UploadFile
from fastapi.responses import JSONResponse
from fastapi.concurrency import run_in_threadpool
from pydantic import BaseModel
from docx import Document
from starlette.responses import Response

from core_models.transcription import get_shared_transcriber
from utils.correction_feedback import save_feedback

router = APIRouter()

_CONTENT_TYPE_EXT = {
    "audio/webm": ".webm",
    "audio/ogg": ".ogg",
    "audio/wav": ".wav",
    "audio/x-wav": ".wav",
    "audio/mp4": ".m4a",
    "audio/mpeg": ".mp3",
    "audio/flac": ".flac",
}

_ALLOWED_EXTS = {".mp3", ".wav", ".m4a", ".ogg", ".flac", ".webm"}
_MAX_UPLOAD_BYTES = 500 * 1024 * 1024


class RectifyRequest(BaseModel):
    original_transcript: str
    edited_transcript: str
    remarks: str | None = None


class ExportDocxRequest(BaseModel):
    transcript: str
    title: str | None = None


def _pick_extension(audio: UploadFile) -> str:
    file_ext = Path(audio.filename or "").suffix.lower()
    if file_ext in _ALLOWED_EXTS:
        return file_ext
    content_type = (audio.content_type or "").lower()
    return _CONTENT_TYPE_EXT.get(content_type, ".webm")


def _is_markdown_table_line(line: str) -> bool:
    raw = (line or "").strip()
    return raw.startswith("|") and raw.endswith("|")


def _is_table_separator_row(line: str) -> bool:
    raw = (line or "").strip().strip("|").strip()
    if not raw:
        return False
    return all(part.strip().replace(":", "").replace("-", "") == "" for part in raw.split("|"))


def _table_cells(line: str) -> List[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _parse_number_token(token: str) -> Optional[int]:
    if not token:
        return None
    token = token.strip().lower()
    if token.isdigit():
        return int(token)
    words = {
        "one": 1,
        "two": 2,
        "three": 3,
        "four": 4,
        "five": 5,
        "six": 6,
        "seven": 7,
        "eight": 8,
        "nine": 9,
        "ten": 10,
    }
    return words.get(token)


def _parse_ordinal_token(token: str) -> Optional[int]:
    if not token:
        return None
    token = token.strip().lower()
    ord_map = {
        "first": 1,
        "second": 2,
        "third": 3,
        "fourth": 4,
        "fifth": 5,
        "sixth": 6,
        "seventh": 7,
        "eighth": 8,
        "ninth": 9,
        "tenth": 10,
    }
    if token in ord_map:
        return ord_map[token]
    m = re.fullmatch(r"(\d+)(?:st|nd|rd|th)?", token)
    if m:
        return int(m.group(1))
    return _parse_number_token(token)


def _cleanup_cell_value(raw: str) -> str:
    out = (raw or "").strip()
    out = re.sub(r"^[\s,.;:\-]+", "", out)
    out = re.sub(r"^(?:is|are|equals?|=)\s+", "", out, flags=re.IGNORECASE)
    out = re.sub(r"^(?:see\s+here|data|value|text)\s+", "", out, flags=re.IGNORECASE)

    # Trim at earliest table/footer terminator phrase (incl. common ASR misspellings).
    terminators = [
        r"\[\[\s*table\s*_?\s*end\s*\]\]",
        r"\[\[\s*table\s*_?\s*(?:row|col(?:umn)?|start)\s*\]\]",
        r"\b(?:and|end|in|n)\s*[-]?\s*table\b",
        r"\b(?:dictated|dictatd|dicteted|dictate)\b",
        r"\b(?:transcribed|transcribe|atranscribed|a\s*transcribed)\b",
        r"\bprinting\b",
        r"\bprinted\s+by\b",
        r"\bjudicial\s+(?:stt|s\s*t\s*t)\s+system\b",
        r"\blegal\s+transcription\s+done\s+by\b",
        r"\bend\s+of\s+transcription\b",
        r"\bopenbooklm\b",
        r"[─\-]{8,}",
    ]
    cut_at = None
    for pat in terminators:
        m = re.search(pat, out, flags=re.IGNORECASE)
        if m:
            cut_at = m.start() if cut_at is None else min(cut_at, m.start())
    if cut_at is not None:
        out = out[:cut_at]

    # Remove any stray table marker token fragments still left in cell text.
    out = re.sub(r"\[\[\s*table\s*_?\s*[a-z]+\s*\]\]", "", out, flags=re.IGNORECASE)
    out = re.sub(r"[\s,.;:\-]+$", "", out)
    return out.strip()


def _extract_command_table(text: str) -> Tuple[Optional[dict], str, str]:
    """
    Parse command-style table instructions such as:
    "make table of two columns and two rows ... data in cell 1 is X ..."
    """
    raw = text or ""
    m = re.search(
        r"\b(?:make|create|generate)\s+table\s+of\s+(\w+)\s+columns?(?:\s+and\s+(\w+)\s+rows?)?",
        raw,
        flags=re.IGNORECASE,
    )
    if not m:
        return None, raw, ""

    cols = _parse_number_token(m.group(1) or "")
    rows = _parse_number_token(m.group(2) or "") if m.group(2) else None
    if not cols or cols < 1:
        return None, raw, ""

    cells = {}
    for cm in re.finditer(
        r"\b(?:data\s+in\s+)?cell\s+(\w+)\s+(?:is|=|:)\s*([^.;\n]+)",
        raw,
        flags=re.IGNORECASE,
    ):
        idx = _parse_number_token(cm.group(1) or "")
        if not idx or idx < 1:
            continue
        cells[idx] = _cleanup_cell_value(cm.group(2))

    # Alternate narration pattern:
    # "first cell ... second cell ... third cell ..."
    marker_re = re.compile(
        r"\b(?:(first|second|third|fourth|fifth|sixth|seventh|eighth|ninth|tenth|\d+(?:st|nd|rd|th)?)\s+cell|cell\s+(\d+))\b",
        flags=re.IGNORECASE,
    )
    markers = list(marker_re.finditer(raw))
    marker_span: Optional[Tuple[int, int]] = None
    post_table_suffix = ""
    last_marker_end = 0
    if markers:
        marker_span = (markers[0].start(), markers[-1].end())
        last_marker_end = markers[-1].end()
        for idx_m, mkr in enumerate(markers):
            token = mkr.group(1) or mkr.group(2) or ""
            idx = _parse_ordinal_token(token)
            if not idx or idx < 1:
                continue
            start = mkr.end()
            end = markers[idx_m + 1].start() if idx_m + 1 < len(markers) else len(raw)
            value_raw = raw[start:end]
            value_for_cell = value_raw
            if idx_m == len(markers) - 1:
                tail_match = re.search(
                    r"\b(?:and|end|in|n)\s*[-]?\s*table\b",
                    value_raw,
                    flags=re.IGNORECASE,
                )
                if tail_match:
                    # Split final cell value from text dictated after table end.
                    term_start = tail_match.start()
                    term_end = tail_match.end()
                    value_for_cell = value_raw[:term_start]
                    post_table_suffix = value_raw[term_end:].strip()
                    last_marker_end = start + term_end
            value = _cleanup_cell_value(value_for_cell)
            if value:
                cells[idx] = value

    if not cells:
        return None, raw, ""

    max_idx = max(cells.keys())
    inferred_rows = (max_idx + cols - 1) // cols
    final_rows = max(1, rows or inferred_rows)

    matrix = []
    cell_idx = 1
    for _ in range(final_rows):
        row = []
        for _ in range(cols):
            row.append(cells.get(cell_idx, ""))
            cell_idx += 1
        matrix.append(row)

    cleaned = raw
    prefix = raw[:m.start()] if m else ""
    suffix_only = ""
    if marker_span:
        # Preserve narrative spoken after table instructions (e.g., "next para ...").
        # Keep only suffix beyond the final cell marker and strip table-ending cues.
        _, e = marker_span
        suffix_from_raw = raw[last_marker_end:] if last_marker_end <= len(raw) else (raw[e:] if e <= len(raw) else "")
        suffix = post_table_suffix or suffix_from_raw
        suffix = re.sub(r"^[\s,.;:\-]+", "", suffix)
        suffix = re.sub(r"^(?:and|end|in|n)\s*[-]?\s*table\b", "", suffix, flags=re.IGNORECASE)
        suffix = re.sub(r"^[\s,.;:\-]+", "", suffix)
        suffix_only = suffix
        prefix = re.sub(r"[\s,.;:\-]+$", "", (prefix or ""))
        cleaned = f"{prefix}\n{suffix}".strip() if prefix and suffix else (prefix or suffix)
    else:
        # Keep dictated content before table command and anything after the command sentence.
        cleaned = re.sub(
            r"\b(?:make|create|generate)\s+table\s+of\s+\w+\s+columns?(?:\s+and\s+\w+\s+rows?)?[^.]*",
            "",
            cleaned,
            flags=re.IGNORECASE,
        )
    cleaned = re.sub(r"\band\s+table\b[\s\W]*$", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s{2,}", " ", cleaned).strip()
    before_text = re.sub(r"\s{2,}", " ", (prefix or "")).strip()
    after_text = suffix_only if marker_span else cleaned
    return {"rows": final_rows, "cols": cols, "data": matrix}, before_text, after_text


def _append_transcript_content(doc: Document, transcript: str) -> None:
    lines = (transcript or "").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if _is_markdown_table_line(line):
            table_lines = []
            while i < len(lines) and _is_markdown_table_line(lines[i].rstrip()):
                table_lines.append(lines[i].rstrip())
                i += 1

            if not table_lines:
                continue

            header = _table_cells(table_lines[0])
            data_rows = table_lines[1:]
            if data_rows and _is_table_separator_row(data_rows[0]):
                data_rows = data_rows[1:]

            col_count = max(1, len(header))
            table = doc.add_table(rows=1, cols=col_count)
            table.style = "Table Grid"

            for c in range(col_count):
                table.cell(0, c).text = header[c] if c < len(header) else ""

            for raw_row in data_rows:
                cells = _table_cells(raw_row)
                row = table.add_row().cells
                for c in range(col_count):
                    row[c].text = cells[c] if c < len(cells) else ""
            continue

        if line.strip():
            trimmed = line.strip()
            if not re.fullmatch(r"[.,;:!?-]+", trimmed):
                doc.add_paragraph(trimmed)
        else:
            doc.add_paragraph("")
        i += 1


def _build_docx_bytes(transcript: str, title: str | None = None) -> bytes:
    doc = Document()
    doc.add_heading((title or "Judicial Transcript").strip(), level=1)

    command_table, before_text, after_text = _extract_command_table(transcript)
    if command_table:
        if before_text:
            _append_transcript_content(doc, before_text)
            doc.add_paragraph("")

        table = doc.add_table(rows=command_table["rows"], cols=command_table["cols"])
        table.style = "Table Grid"
        for r in range(command_table["rows"]):
            for c in range(command_table["cols"]):
                table.cell(r, c).text = command_table["data"][r][c]
        doc.add_paragraph("")

        if after_text:
            _append_transcript_content(doc, after_text)
    else:
        _append_transcript_content(doc, transcript)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


@router.post("/transcribe")
async def transcribe_audio(
    audio: UploadFile = File(...),
    language: str = "en",
    format_type: str = "high_court",
):
    """
    Transcribe uploaded audio file.
    """
    ext = _pick_extension(audio)
    if ext not in _ALLOWED_EXTS:
        raise HTTPException(status_code=400, detail="Unsupported audio format")

    raw = await audio.read()
    if not raw:
        raise HTTPException(status_code=400, detail="Empty audio upload")
    if len(raw) > _MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=400, detail="Audio file too large")

    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp.write(raw)
            temp_path = tmp.name

        # Remark: run CPU-heavy transcription in worker thread so health/ws stays responsive.
        result = await run_in_threadpool(
            get_shared_transcriber().transcribe_file,
            temp_path,
            language,
            format_type,
            (audio.filename or temp_path),
        )

        response = {
            "success": True,
            "transcript": result.get("text", ""),
            "formatted": result.get("formatted", ""),
            "format_type": format_type,
            "language": result.get("language", language),
            "confidence": result.get("confidence", 0),
            "duration": result.get("duration", 0),
            "timestamp": datetime.now().isoformat(),
        }
        return JSONResponse(content=response)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Transcription failed: {str(e)}")
    finally:
        if temp_path:
            try:
                Path(temp_path).unlink(missing_ok=True)
            except Exception:
                pass


@router.get("/formats")
async def get_supported_formats():
    """Get supported formats and languages."""
    return {
        "supported_formats": [
            {"id": "high_court", "name": "High Court Judgment"},
            {"id": "supreme_court", "name": "Supreme Court Judgment"},
            {"id": "district_court", "name": "District Court Order"},
        ],
        "supported_languages": [
            {"code": "en", "name": "English"},
            {"code": "hi", "name": "Hindi"},
        ],
        "max_file_size_mb": 500,
        "supported_audio_formats": [".mp3", ".wav", ".m4a", ".ogg", ".flac", ".webm"],
    }


@router.post("/rectify")
async def rectify_transcript(data: RectifyRequest):
    """
    Save user-corrected transcript and apply spoken-formatting commands.
    """
    original = (data.original_transcript or "").strip()
    edited = (data.edited_transcript or "").strip()
    if not original:
        raise HTTPException(status_code=400, detail="original_transcript is required")
    if not edited:
        raise HTTPException(status_code=400, detail="edited_transcript is required")

    transcriber = get_shared_transcriber()
    normalized = transcriber.legal_formatter.normalize_user_edit(edited)
    feedback_info = save_feedback(
        original_transcript=original,
        edited_transcript=normalized,
        remarks=(data.remarks or "").strip() or None,
    )
    # Apply newly submitted corrections to future transcripts without restart.
    transcriber.legal_formatter.reload_feedback_phrase_map()

    return {
        "success": True,
        "corrected_transcript": normalized,
        "remarks_saved": bool((data.remarks or "").strip()),
        "replacement_pairs_detected": len(feedback_info.get("replacement_pairs") or []),
    }


@router.get("/rectify/guide")
async def get_rectify_guide():
    return {
        "title": "STT Rectification Guide",
        "steps": [
            "Run transcription",
            "Review and edit transcript",
            "Add optional remarks",
            "Submit rectification",
            "Use corrected transcript for final output",
        ],
        "spoken_commands": [
            "open bracket / in bracket -> (",
            "close bracket -> )",
            "comma -> ,",
            "full stop -> .",
            "question mark -> ?",
            "next para / next paragraph -> new paragraph",
            "start table, next column, next row, end table -> markdown table",
        ],
    }


@router.post("/export/docx")
async def export_transcript_docx(data: ExportDocxRequest):
    transcript = (data.transcript or "").strip()
    if not transcript:
        raise HTTPException(status_code=400, detail="transcript is required")

    payload = _build_docx_bytes(transcript=transcript, title=data.title)
    safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", (data.title or "legal_transcript")).strip("_")
    if not safe_name:
        safe_name = "legal_transcript"

    return Response(
        content=payload,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{safe_name}.docx"',
        },
    )
